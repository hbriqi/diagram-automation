from diagrams import Cluster, Diagram, Node
from diagrams.k8s.compute import Pod
from diagrams.onprem import compute, container, database, client, inmemory, queue
from diagrams.azure import integration as azure_integration, identity as azure_identity, web as azure_web
from diagrams.aws import network as aws_network
from diagrams.aws import security as aws_security
from diagrams.generic import device

from docx import Document
from docx.shared import Inches

import datetime

def BuildModel(components, flatlist):
    for component in components:
        if component["type"] == "cluster":
            with Cluster(component["name"], direction="LR") as cluster:
                BuildModel(component["components"], flatlist)
                flatlist[component["name"]] = cluster 

        elif component["type"] == "pc-client":
            flatlist[component["name"]] = client.Client(component["name"])

        elif component["type"] == "mobile-client":
            flatlist[component["name"]] = device.Mobile(component["name"])
        
        elif component["type"] == "dmz-fw":
            flatlist[component["name"]] = aws_security.FirewallManager(component["name"])
        
        elif component["type"] == "dmz-lb":
            flatlist[component["name"]] = aws_network.ElbNetworkLoadBalancer(component["name"])
        
        elif component["type"] == "api-manager":
            flatlist[component["name"]] = azure_integration.APIManagement(component["name"]) #aws_network.APIGateway(component["name"])
        
        elif component["type"] == "iam-manager":
            flatlist[component["name"]] = azure_identity.ADPrivilegedIdentityManagement(component["name"])
        
        elif component["type"] == "content":
            flatlist[component["name"]] = aws_network.NetworkingAndContentDelivery(component["name"])
        
        elif component["type"] == "bs-container":
            flatlist[component["name"]] = Pod(component["name"])
        
        elif component["type"] == "ms-sql-db":
            flatlist[component["name"]] = database.Mssql(component["name"])
        
        elif component["type"] == "pg-sql-db":
            flatlist[component["name"]] = database.Postgresql(component["name"])
        
        elif component["type"] == "redis-db":
           flatlist[component["name"]] = inmemory.Redis(component["name"])
        
        elif component["type"] == "int-service":
            flatlist[component["name"]] = azure_web.AppServices(component["name"])

        elif component["type"] == "queue":
            flatlist[component["name"]] = queue.Rabbitmq(component["name"])


def LinkModel(connections , flatlist):
    for item in connections:
        flatlist[item["from"]] >> flatlist[item["to"]]


def CreateWordDocument(data):
    document = Document()

    #doucment title
    document.add_heading(data["project_name"], 0)
    document.add_page_break()

    #document info
    document.add_heading('Document Info', level=1)
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Project Name'
    hdr_cells[1].text = 'Date'
    hdr_cells[2].text = 'Description'
    row_cells = table.add_row().cells
    row_cells[0].text = "SAIP IPN"
    date = datetime.datetime.now()
    row_cells[1].text = str(date.day) + "/" + str(date.month) +"/" + str(date.year)
    row_cells[2].text = ""
    
    #physical diagram
    document.add_heading('Physical Diagram', level=1)
    picture_file = data["diagram_name"].replace(" ", "_") +'.png';
    document.add_picture(picture_file, width=Inches(7.0))
    document.add_page_break()
    document.save(data["project_name"] + '.docx')



    